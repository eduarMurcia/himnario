from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from pypdf import PdfReader

from hymn_studio.models import Lyrics, Stanza

_CHORUS_KEYWORDS = ("coro", "estribillo", "chorus")


class LyricsExtractionError(RuntimeError):
    pass


class LyricsExtractor:
    """Extracts hymn lyrics from a PDF or Word document into title + stanzas."""

    def __init__(
        self, max_lines_per_stanza: int = 4, repeat_chorus_after_verses: bool = True
    ) -> None:
        """max_lines_per_stanza: stanzas longer than this are split into exactly two
        slides so the rendered text stays legible instead of being shrunk to fit. An
        even line count splits into two equal halves; an odd count gives the first
        half the extra line (5 -> 3+2, 7 -> 4+3).

        repeat_chorus_after_verses: hymn sheets usually write the chorus once to save
        space, but it is sung after every verse. When True (default), the chorus is
        removed from its original position and re-inserted after each verse."""
        self._max_lines_per_stanza = max_lines_per_stanza
        self._repeat_chorus_after_verses = repeat_chorus_after_verses

    def extract(self, path: Path) -> Lyrics:
        suffix = path.suffix.lower()
        if suffix == ".pdf":
            return self.extract_pdf(path)
        if suffix == ".docx":
            return self.extract_docx(path)
        raise LyricsExtractionError(f"Unsupported lyrics file type: {suffix}")

    def extract_pdf(self, path: Path) -> Lyrics:
        return self._parse(path.stem, self._read_pdf(path))

    def extract_docx(self, path: Path) -> Lyrics:
        return self._parse(path.stem, self._read_docx(path))

    def _read_pdf(self, path: Path) -> str:
        try:
            reader = PdfReader(str(path))
        except Exception as error:
            raise LyricsExtractionError(f"Could not read PDF: {path}") from error

        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(pages)

    def _read_docx(self, path: Path) -> str:
        try:
            document = Document(str(path))
        except Exception as error:
            raise LyricsExtractionError(f"Could not read Word document: {path}") from error

        return "\n".join(paragraph.text for paragraph in document.paragraphs)

    def _parse(self, default_title: str, raw_text: str) -> Lyrics:
        lines = [line.strip() for line in raw_text.splitlines()]

        first_content_index = next((i for i, line in enumerate(lines) if line), None)
        if first_content_index is None:
            raise LyricsExtractionError("No lyrics content found.")

        title = lines[first_content_index] or default_title
        return self.build_lyrics(title, lines[first_content_index + 1 :])

    def build_lyrics(self, title: str, lines: list[str]) -> Lyrics:
        """Builds a Lyrics from a title and the raw (unstripped-of-blanks) lines that
        follow it. Shared by single-hymn parsing and MultiHymnExtractor."""
        blocks = self._split_into_blocks(lines)
        if not blocks:
            raise LyricsExtractionError(f"No stanzas found for '{title}'.")

        stanzas = [self._to_stanza(block) for block in blocks]
        if self._repeat_chorus_after_verses:
            stanzas = self._repeat_chorus(stanzas)
        stanzas = [split for stanza in stanzas for split in self._split_long_stanza(stanza)]
        return Lyrics(title=title, stanzas=stanzas)

    def _split_into_blocks(self, lines: list[str]) -> list[list[str]]:
        blocks: list[list[str]] = []
        current: list[str] = []
        for line in lines:
            if line:
                current.append(line)
            elif current:
                blocks.append(current)
                current = []
        if current:
            blocks.append(current)
        return blocks

    def _to_stanza(self, block: list[str]) -> Stanza:
        is_chorus = self._is_chorus_heading(block[0])
        text_lines = block[1:] if is_chorus and len(block) > 1 else block
        return Stanza(text="\n".join(text_lines), is_chorus=is_chorus)

    def _repeat_chorus(self, stanzas: list[Stanza]) -> list[Stanza]:
        chorus = [stanza for stanza in stanzas if stanza.is_chorus]
        if not chorus:
            return stanzas

        result: list[Stanza] = []
        for stanza in stanzas:
            if stanza.is_chorus:
                continue
            result.append(stanza)
            result.extend(chorus)
        return result

    def _split_long_stanza(self, stanza: Stanza) -> list[Stanza]:
        lines = stanza.text.splitlines()
        if len(lines) <= self._max_lines_per_stanza:
            return [stanza]

        first_size = (len(lines) + 1) // 2
        return [
            Stanza(text="\n".join(lines[:first_size]), is_chorus=stanza.is_chorus),
            Stanza(text="\n".join(lines[first_size:]), is_chorus=stanza.is_chorus),
        ]

    def _is_chorus_heading(self, line: str) -> bool:
        normalized = re.sub(r"[^a-záéíóúñ]", "", line.lower())
        return any(normalized.startswith(keyword) for keyword in _CHORUS_KEYWORDS)


class MultiHymnExtractor:
    """Splits a single Word document containing many hymns into one Lyrics per hymn.

    Each hymn is expected to start with a "Heading 1" styled paragraph (its title),
    followed by Normal-styled verses/chorus in the same format single-hymn documents
    use. Headings with no lyric content before the next heading (e.g. a table-of-
    contents entry) are skipped rather than treated as empty hymns.
    """

    def __init__(self, lyrics_extractor: LyricsExtractor | None = None) -> None:
        self._lyrics_extractor = lyrics_extractor or LyricsExtractor()

    def extract_docx(self, path: Path) -> list[Lyrics]:
        try:
            document = Document(str(path))
        except Exception as error:
            raise LyricsExtractionError(f"Could not read Word document: {path}") from error

        paragraphs = [
            (paragraph.text.strip(), self._is_heading(paragraph))
            for paragraph in document.paragraphs
        ]
        heading_indices = [i for i, (_, is_heading) in enumerate(paragraphs) if is_heading]

        hymns: list[Lyrics] = []
        for position, index in enumerate(heading_indices):
            title = paragraphs[index][0]
            next_heading = heading_indices[position + 1 :]
            end = next_heading[0] if next_heading else len(paragraphs)
            lines = [text for text, _ in paragraphs[index + 1 : end]]

            if not any(line for line in lines):
                continue

            hymns.append(self._lyrics_extractor.build_lyrics(title, lines))

        return hymns

    def _is_heading(self, paragraph) -> bool:
        style_name = paragraph.style.name if paragraph.style else ""
        return style_name.startswith("Heading")
