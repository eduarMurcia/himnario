import sys
import tempfile
import unittest
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "src"))

from docx import Document

from hymn_studio.services.lyrics import LyricsExtractor, MultiHymnExtractor


class MultiHymnExtractorTest(unittest.TestCase):
    def test_splits_hymns_by_heading_and_skips_empty_index(self) -> None:
        with self._temp_docx() as docx_path:
            hymns = MultiHymnExtractor().extract_docx(docx_path)

        self.assertEqual(len(hymns), 2)

        self.assertEqual(hymns[0].title, "A Combatir")
        self.assertEqual(len(hymns[0].stanzas), 2)
        self.assertIn("Resuena la potente voz", hymns[0].stanzas[0].text)
        self.assertTrue(hymns[0].stanzas[1].is_chorus)

        self.assertEqual(hymns[1].title, "Dad loor al Senor")
        self.assertEqual(len(hymns[1].stanzas), 1)

    def test_uses_custom_lyrics_extractor_settings(self) -> None:
        with self._temp_docx() as docx_path:
            extractor = LyricsExtractor(max_lines_per_stanza=1, repeat_chorus_after_verses=False)
            hymns = MultiHymnExtractor(extractor).extract_docx(docx_path)

        # verse1 has 2 lines, split into 2 slides of 1 line each with chorus repeat off
        self.assertEqual(hymns[0].title, "A Combatir")
        self.assertEqual(len(hymns[0].stanzas), 3)

    def _temp_docx(self):
        document = Document()
        document.add_paragraph("Indice de Contenido", style="Heading 1")

        document.add_paragraph("A Combatir", style="Heading 1")
        document.add_paragraph("Resuena la potente voz,")
        document.add_paragraph("del buen Jesus que llamando esta.")
        document.add_paragraph("")
        document.add_paragraph("Coro:")
        document.add_paragraph("A la batalla, oh cristiano.")

        document.add_paragraph("Dad loor al Senor", style="Heading 1")
        document.add_paragraph("A Dios demos gloria.")

        temp_dir = tempfile.TemporaryDirectory()
        path = Path(temp_dir.name) / "himnario.docx"
        document.save(str(path))

        class _Ctx:
            def __enter__(self_inner):
                return path

            def __exit__(self_inner, *exc_info):
                temp_dir.cleanup()

        return _Ctx()


if __name__ == "__main__":
    unittest.main()
