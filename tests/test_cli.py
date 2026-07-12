import sys
import tempfile
import unittest
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "src"))

from docx import Document
from PIL import Image

from hymn_studio.cli import main
from hymn_studio.project import ProjectRepository

WINDOWS_ARIAL = Path("C:/Windows/Fonts/arial.ttf")


@unittest.skipUnless(WINDOWS_ARIAL.exists(), "Requires a system TTF font (arial.ttf).")
class CliBuildCommandTest(unittest.TestCase):
    def test_build_with_pillow_engine_creates_project(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)

            lyrics_path = self._make_lyrics_docx(temp_path / "hymn.docx")
            audio_path = temp_path / "audio.mp3"
            audio_path.write_bytes(b"fake-audio")
            background_path = temp_path / "bg.png"
            Image.new("RGB", (400, 300), color=(0, 0, 0)).save(background_path)
            out_dir = temp_path / "project"

            exit_code = main(
                [
                    "build",
                    str(lyrics_path),
                    str(audio_path),
                    "--out",
                    str(out_dir),
                    "--engine",
                    "pillow",
                    "--cover-background",
                    str(background_path),
                    "--stanza-background",
                    str(background_path),
                    "--font",
                    str(WINDOWS_ARIAL),
                ]
            )

            self.assertEqual(exit_code, 0)

            project_path = out_dir / "Test Hymn.hymn"
            self.assertTrue(project_path.exists())

            project = ProjectRepository().load(project_path)
            self.assertEqual(project.audio_path, audio_path)
            self.assertEqual(project.image_folder, out_dir / "slides")
            self.assertEqual(project.timestamps, [])

            slide_files = sorted((out_dir / "slides").glob("*.png"))
            self.assertEqual(len(slide_files), 3)

    def _make_lyrics_docx(self, path: Path) -> Path:
        document = Document()
        document.add_paragraph("Test Hymn")
        document.add_paragraph("")
        document.add_paragraph("Verse one line")
        document.add_paragraph("")
        document.add_paragraph("Coro")
        document.add_paragraph("Chorus line")
        document.save(str(path))
        return path


@unittest.skipUnless(WINDOWS_ARIAL.exists(), "Requires a system TTF font (arial.ttf).")
class CliBuildAllCommandTest(unittest.TestCase):
    def test_build_all_matches_audio_and_reports_skipped(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)

            lyrics_path = self._make_multi_hymn_docx(temp_path / "himnario.docx")
            audio_dir = temp_path / "audio"
            audio_dir.mkdir()
            (audio_dir / "Primer Himno.mp3").write_bytes(b"fake-audio")
            # "Segundo Himno" intentionally has no matching audio file.

            background_path = temp_path / "bg.png"
            Image.new("RGB", (400, 300), color=(0, 0, 0)).save(background_path)
            out_dir = temp_path / "out"

            exit_code = main(
                [
                    "build-all",
                    str(lyrics_path),
                    "--audio-dir",
                    str(audio_dir),
                    "--out",
                    str(out_dir),
                    "--engine",
                    "pillow",
                    "--cover-background",
                    str(background_path),
                    "--stanza-background",
                    str(background_path),
                    "--font",
                    str(WINDOWS_ARIAL),
                ]
            )

            self.assertEqual(exit_code, 0)

            project_path = out_dir / "Primer Himno" / "Primer Himno.hymn"
            self.assertTrue(project_path.exists())
            self.assertFalse((out_dir / "Segundo Himno").exists())

    def _make_multi_hymn_docx(self, path: Path) -> Path:
        document = Document()
        document.add_paragraph("Indice", style="Heading 1")

        document.add_paragraph("Primer Himno", style="Heading 1")
        document.add_paragraph("Primera linea del primer himno.")

        document.add_paragraph("Segundo Himno", style="Heading 1")
        document.add_paragraph("Primera linea del segundo himno.")

        document.save(str(path))
        return path


if __name__ == "__main__":
    unittest.main()
