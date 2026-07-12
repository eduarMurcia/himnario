import sys
import unittest
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "src"))

from docx import Document

from hymn_studio.services.lyrics import LyricsExtractionError, LyricsExtractor


class LyricsExtractorTest(unittest.TestCase):
    def test_parses_title_stanzas_and_chorus(self) -> None:
        raw_text = (
            "Grande Es Tu Fidelidad\n"
            "\n"
            "Oh Dios eterno, tu fidelidad\n"
            "es como el sol que nunca dejara\n"
            "\n"
            "Coro:\n"
            "Grande es tu fidelidad\n"
            "grande es tu fidelidad\n"
            "\n"
            "Cada manana nuevas mercedes\n"
            "veo brotar en derredor\n"
        )

        lyrics = LyricsExtractor(repeat_chorus_after_verses=False)._parse(
            "fallback-title", raw_text
        )

        self.assertEqual(lyrics.title, "Grande Es Tu Fidelidad")
        self.assertEqual(len(lyrics.stanzas), 3)

        self.assertFalse(lyrics.stanzas[0].is_chorus)
        self.assertIn("Oh Dios eterno", lyrics.stanzas[0].text)

        self.assertTrue(lyrics.stanzas[1].is_chorus)
        self.assertNotIn("Coro", lyrics.stanzas[1].text)
        self.assertIn("Grande es tu fidelidad", lyrics.stanzas[1].text)

        self.assertFalse(lyrics.stanzas[2].is_chorus)

    def test_repeats_chorus_after_every_verse_by_default(self) -> None:
        raw_text = (
            "Grande Es Tu Fidelidad\n"
            "\n"
            "Verse one\n"
            "\n"
            "Coro:\n"
            "Chorus line\n"
            "\n"
            "Verse two\n"
            "\n"
            "Verse three\n"
        )

        lyrics = LyricsExtractor()._parse("fallback-title", raw_text)

        self.assertEqual(len(lyrics.stanzas), 6)
        expected_chorus_flags = [False, True, False, True, False, True]
        self.assertEqual([s.is_chorus for s in lyrics.stanzas], expected_chorus_flags)
        self.assertEqual(lyrics.stanzas[0].text, "Verse one")
        self.assertEqual(lyrics.stanzas[1].text, "Chorus line")
        self.assertEqual(lyrics.stanzas[2].text, "Verse two")
        self.assertEqual(lyrics.stanzas[3].text, "Chorus line")
        self.assertEqual(lyrics.stanzas[4].text, "Verse three")
        self.assertEqual(lyrics.stanzas[5].text, "Chorus line")

    def test_repeat_chorus_can_be_disabled(self) -> None:
        raw_text = "Title\n\nVerse one\n\nCoro:\nChorus line\n\nVerse two\n"

        lyrics = LyricsExtractor(repeat_chorus_after_verses=False)._parse(
            "fallback-title", raw_text
        )

        self.assertEqual(len(lyrics.stanzas), 3)

    def test_title_is_always_the_first_line_even_without_a_blank_line_after(self) -> None:
        raw_text = "Title Line\nFirst stanza line one\nFirst stanza line two\n"

        lyrics = LyricsExtractor()._parse("fallback-title", raw_text)

        self.assertEqual(lyrics.title, "Title Line")
        self.assertEqual(len(lyrics.stanzas), 1)
        self.assertEqual(lyrics.stanzas[0].text, "First stanza line one\nFirst stanza line two")

    def test_splits_eight_lines_into_equal_halves(self) -> None:
        raw_text = (
            "Dad loor al Senor\n"
            "linea 1\n"
            "linea 2\n"
            "linea 3\n"
            "linea 4\n"
            "linea 5\n"
            "linea 6\n"
            "linea 7\n"
            "linea 8\n"
        )

        lyrics = LyricsExtractor(max_lines_per_stanza=4)._parse("fallback-title", raw_text)

        self.assertEqual(len(lyrics.stanzas), 2)
        self.assertEqual(lyrics.stanzas[0].text, "linea 1\nlinea 2\nlinea 3\nlinea 4")
        self.assertEqual(lyrics.stanzas[1].text, "linea 5\nlinea 6\nlinea 7\nlinea 8")

    def test_does_not_split_short_stanzas(self) -> None:
        raw_text = "Title\nlinea 1\nlinea 2\nlinea 3\n"

        lyrics = LyricsExtractor(max_lines_per_stanza=4)._parse("fallback-title", raw_text)

        self.assertEqual(len(lyrics.stanzas), 1)

    def test_splits_odd_stanza_with_larger_first_half(self) -> None:
        raw_text = "Title\n" + "\n".join(f"linea {i}" for i in range(1, 6))

        lyrics = LyricsExtractor(max_lines_per_stanza=4)._parse("fallback-title", raw_text)

        self.assertEqual(len(lyrics.stanzas), 2)
        self.assertEqual(lyrics.stanzas[0].text, "linea 1\nlinea 2\nlinea 3")
        self.assertEqual(lyrics.stanzas[1].text, "linea 4\nlinea 5")

    def test_splits_seven_lines_into_four_and_three(self) -> None:
        raw_text = "Title\n" + "\n".join(f"linea {i}" for i in range(1, 8))

        lyrics = LyricsExtractor(max_lines_per_stanza=4)._parse("fallback-title", raw_text)

        self.assertEqual(len(lyrics.stanzas), 2)
        self.assertEqual(lyrics.stanzas[0].text, "linea 1\nlinea 2\nlinea 3\nlinea 4")
        self.assertEqual(lyrics.stanzas[1].text, "linea 5\nlinea 6\nlinea 7")

    def test_splits_six_lines_into_equal_halves(self) -> None:
        raw_text = "Title\n" + "\n".join(f"linea {i}" for i in range(1, 7))

        lyrics = LyricsExtractor(max_lines_per_stanza=4)._parse("fallback-title", raw_text)

        self.assertEqual(len(lyrics.stanzas), 2)
        self.assertEqual(lyrics.stanzas[0].text, "linea 1\nlinea 2\nlinea 3")
        self.assertEqual(lyrics.stanzas[1].text, "linea 4\nlinea 5\nlinea 6")

    def test_split_preserves_chorus_flag_on_every_part(self) -> None:
        raw_text = (
            "Title\n"
            "verse\n"
            "\n"
            "Coro\n"
            "c1\nc2\nc3\nc4\nc5\nc6\n"
        )

        lyrics = LyricsExtractor(max_lines_per_stanza=4)._parse("fallback-title", raw_text)

        chorus_parts = [s for s in lyrics.stanzas if s.is_chorus]
        self.assertEqual(len(chorus_parts), 2)

    def test_raises_on_empty_content(self) -> None:
        with self.assertRaises(LyricsExtractionError):
            LyricsExtractor()._parse("fallback-title", "   \n\n  ")

    def test_extracts_from_docx(self) -> None:
        with self._temp_docx() as docx_path:
            lyrics = LyricsExtractor().extract_docx(docx_path)

        self.assertEqual(lyrics.title, "Test Hymn")
        self.assertEqual(len(lyrics.stanzas), 2)
        self.assertTrue(lyrics.stanzas[1].is_chorus)

    def _temp_docx(self):
        import tempfile

        document = Document()
        document.add_paragraph("Test Hymn")
        document.add_paragraph("")
        document.add_paragraph("First line of verse one")
        document.add_paragraph("Second line of verse one")
        document.add_paragraph("")
        document.add_paragraph("Estribillo")
        document.add_paragraph("Chorus line one")

        temp_dir = tempfile.TemporaryDirectory()
        path = Path(temp_dir.name) / "hymn.docx"
        document.save(str(path))

        class _Ctx:
            def __enter__(self_inner):
                return path

            def __exit__(self_inner, *exc_info):
                temp_dir.cleanup()

        return _Ctx()


if __name__ == "__main__":
    unittest.main()
